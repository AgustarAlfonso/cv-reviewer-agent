"""
Module for parsing PDF files to extract text.
Uses PyMuPDF (fitz) for text extraction.
"""
import fitz  # PyMuPDF
import io
from docx import Document

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extracts text from a PDF file provided as bytes.
    
    Args:
        file_bytes (bytes): The contents of the PDF file.
        
    Returns:
        str: The extracted text from the entire PDF.
    """
    text = ""
    try:
        # Open the PDF from memory
        pdf_document = fitz.open(stream=file_bytes, filetype="pdf")
        
        # Iterate through all pages and extract text
        for page_num in range(pdf_document.page_count):
            page = pdf_document.load_page(page_num)
            text += page.get_text("text") + "\n"
            
        pdf_document.close()
        return text.strip()
    except Exception as e:
        # In a real app, you might want to log this error
        raise ValueError(f"Failed to parse PDF file: {str(e)}")

def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extracts text from a DOCX file provided as bytes.
    Extracts both paragraphs and tables (since CVs often use tables for layout).
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        lines = []
        
        # Extract paragraph text
        for p in doc.paragraphs:
            if p.text.strip():
                lines.append(p.text.strip())
                
        # Extract table text
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_data.append(cell.text.strip())
                if row_data:
                    lines.append(" | ".join(row_data))
                    
        return "\n".join(lines).strip()
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX file: {str(e)}")
