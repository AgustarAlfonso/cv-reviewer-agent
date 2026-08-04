import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from schemas import StructuredCV

def set_font(run, size=11, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def add_horizontal_line(paragraph):
    # This adds a bottom border to a paragraph, simulating a horizontal line
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)

def add_bullet_paragraph(document, text, size=11):
    """
    Adds a real bullet point: a hanging indent with an explicit bullet glyph,
    so it's visually a list item (not just an indented paragraph) while
    remaining plain text for ATS parsers (no numbering.xml dependency).
    """
    bp = document.add_paragraph()
    bp.paragraph_format.space_after = Pt(0)
    bp.paragraph_format.left_indent = Inches(0.5)
    bp.paragraph_format.first_line_indent = Inches(-0.25)
    marker_run = bp.add_run("•\t")
    set_font(marker_run, size=size)
    text_run = bp.add_run(text)
    set_font(text_run, size=size)
    return bp

def keep_with_next(paragraph):
    """
    Marks a paragraph to stay on the same page as the paragraph that follows it,
    preventing a heading/title from being orphaned at the bottom of a page while
    its content flows to the next page.
    """
    paragraph.paragraph_format.keep_with_next = True

def create_docx_from_structured_cv(cv_data: StructuredCV) -> io.BytesIO:
    """
    Creates an ATS-friendly single-column DOCX file from the StructuredCV data,
    matching the user's specific template formatting perfectly.
    """
    document = Document()

    # Set document margins
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # --- Header (Name and Contact) ---
    name_paragraph = document.add_paragraph()
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_paragraph.add_run(cv_data.header.name.upper())
    set_font(name_run, size=14, bold=True)

    headline_paragraph = document.add_paragraph()
    headline_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headline_run = headline_paragraph.add_run(cv_data.header.headline)
    set_font(headline_run, size=11, italic=True)

    contact_p1 = document.add_paragraph()
    contact_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p1_run = contact_p1.add_run(f"{cv_data.header.location}  |  {cv_data.header.phone}  |  {cv_data.header.email}")
    set_font(contact_p1_run, size=11)

    contact_p1.paragraph_format.space_after = Pt(0)
    headline_paragraph.paragraph_format.space_after = Pt(0)
    name_paragraph.paragraph_format.space_after = Pt(0)

    keep_with_next(name_paragraph)
    keep_with_next(headline_paragraph)
    keep_with_next(contact_p1)

    contact_p2 = document.add_paragraph()
    contact_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p2_run = contact_p2.add_run(f"Portfolio: {cv_data.header.portfolio}  |  LinkedIn: {cv_data.header.linkedin}  |  GitHub: {cv_data.header.github}")
    set_font(contact_p2_run, size=11)

    # Helper function for section headings
    def add_section_heading(text: str):
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_font(run, size=12, bold=True)
        add_horizontal_line(p)
        keep_with_next(p)  # a heading should never end up alone at the bottom of a page
        return p

    # --- Professional Summary ---
    if cv_data.summary:
        add_section_heading("Summary")
        summary_p = document.add_paragraph()
        summary_run = summary_p.add_run(cv_data.summary)
        set_font(summary_run, size=11)

    # --- Technical Skills ---
    if cv_data.skills:
        add_section_heading("Technical Skills")
        for skill_cat in cv_data.skills:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            cat_run = p.add_run(skill_cat.category + ": ")
            set_font(cat_run, size=11, bold=True)
            items_run = p.add_run(skill_cat.items)
            set_font(items_run, size=11)

    # --- Experience ---
    if cv_data.experience:
        add_section_heading("Experience")
        for exp in cv_data.experience:
            role_p = document.add_paragraph()
            role_p.paragraph_format.space_after = Pt(0)
            role_run = role_p.add_run(exp.role)
            set_font(role_run, size=11, bold=True)
            keep_with_next(role_p)

            company_p = document.add_paragraph()
            company_p.paragraph_format.space_after = Pt(4)
            comp_run = company_p.add_run(f"{exp.company}  |  {exp.duration}  |  {exp.location}")
            set_font(comp_run, size=11, italic=True)
            if exp.bullets:
                keep_with_next(company_p)  # keep attached to its first bullet

            for i, bullet in enumerate(exp.bullets):
                bp = add_bullet_paragraph(document, bullet)
                if i < len(exp.bullets) - 1:
                    keep_with_next(bp)  # keep a role's bullets from splitting mid-list

    # --- Projects ---
    if cv_data.projects:
        add_section_heading("Projects")
        for proj in cv_data.projects:
            name_p = document.add_paragraph()
            name_p.paragraph_format.space_after = Pt(0)
            name_p.paragraph_format.space_before = Pt(6)
            name_run = name_p.add_run(proj.name)
            set_font(name_run, size=11, bold=True)
            keep_with_next(name_p)

            tech_p = document.add_paragraph()
            tech_p.paragraph_format.space_after = Pt(4)
            tech_run = tech_p.add_run(f"{proj.context}  |  {proj.technologies}  |  {proj.link}")
            set_font(tech_run, size=11, italic=True)
            if proj.bullets:
                keep_with_next(tech_p)  # keep attached to its first bullet

            for i, bullet in enumerate(proj.bullets):
                bp = add_bullet_paragraph(document, bullet)
                if i < len(proj.bullets) - 1:
                    keep_with_next(bp)

    # --- Publication ---
    if cv_data.publications:
        add_section_heading("Publication")
        for pub in cv_data.publications:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(pub.citation)
            set_font(run, size=11)

    # --- Certifications ---
    if cv_data.certifications:
        add_section_heading("Certifications")
        for cert in cv_data.certifications:
            add_bullet_paragraph(document, f"{cert.name} - {cert.issuer}")

    # --- Education ---
    if cv_data.education:
        add_section_heading("Education")
        for edu in cv_data.education:
            inst_p = document.add_paragraph()
            inst_p.paragraph_format.space_after = Pt(0)
            inst_p.paragraph_format.space_before = Pt(6)
            inst_run = inst_p.add_run(edu.institution)
            set_font(inst_run, size=11, bold=True)
            keep_with_next(inst_p)

            deg_p = document.add_paragraph()
            deg_p.paragraph_format.space_after = Pt(4)
            deg_run = deg_p.add_run(f"{edu.degree}  |  {edu.duration}  |  {edu.gpa}")
            set_font(deg_run, size=11)

    # --- Languages & Other Skills ---
    if cv_data.languages_other:
        add_section_heading("Languages & Other Skills")

        lang_p = document.add_paragraph()
        lang_p.paragraph_format.space_after = Pt(0)
        lang_title = lang_p.add_run("Languages: ")
        set_font(lang_title, size=11, bold=True)
        lang_text = lang_p.add_run(cv_data.languages_other.languages)
        set_font(lang_text, size=11)

        other_p = document.add_paragraph()
        other_title = other_p.add_run("Other Skills: ")
        set_font(other_title, size=11, bold=True)
        other_text = other_p.add_run(cv_data.languages_other.other_skills)
        set_font(other_text, size=11)

    # Save to BytesIO stream
    doc_stream = io.BytesIO()
    document.save(doc_stream)
    doc_stream.seek(0)
    return doc_stream